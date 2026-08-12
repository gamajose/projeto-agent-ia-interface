from __future__ import annotations

import io
import re
from copy import deepcopy
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.services.n2_documentation import sanitize_n2_review


NAVY = "0B1F44"
BLUE = "2672D5"
LIGHT_BLUE = "DCE9FC"
RED = "A6242E"
DARK = "10151D"

SUMMARY_SECTIONS = (
    "Introdução", "Infraestrutura", "Inventário", "Informações de Infraestrutura do Inventário",
    "Banco de Dados", "Informações para Ativação na TOTVS", "Estrutura SGDB", "TNSNAMES",
    "Mapeamento do Winthor", "Política de Backup", "Database Oracle", "Backup do Sistema ERP",
    "Métodos de Execução e Validações", "Retenção de Backups Obsoletos", "Dados Redundância",
    "Replicação standby", "Monitoramento", "Considerações Finais",
)

SENSITIVE_KEYS = {
    "password", "senha", "secret", "token", "community", "private_key", "credential",
    "compartilhamento_senha", "monitoring_password", "database_password", "winthor_password",
}


def _text(value: Any) -> str:
    return str(value or "").strip()


def _clean_filename(value: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-.")
    return safe or "documentacao-n2"


def _field_map(review: dict[str, Any]) -> dict[str, str]:
    values: dict[str, str] = {}
    for section in review.get("sections") or []:
        if not isinstance(section, dict):
            continue
        for field in section.get("fields") or []:
            if not isinstance(field, dict):
                continue
            key = _text(field.get("key"))
            if not key or any(word in key.casefold() for word in SENSITIVE_KEYS):
                continue
            values[key] = _text(field.get("value"))
    return values


def _font(size: int, *, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _fit(value: str, width: int, font) -> str:
    text = _text(value)
    if not text:
        return "-"
    probe = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(probe)
    if draw.textlength(text, font=font) <= width:
        return text
    suffix = "..."
    while text and draw.textlength(text + suffix, font=font) > width:
        text = text[:-1]
    return text + suffix


def build_cover_png(review: dict[str, Any]) -> bytes:
    """Gera capa inspirada no template 2Com enviado, sem depender de asset externo."""
    width, height = 827, 1169
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    top = (6, 28, 71)
    bottom = (235, 240, 248)
    for y in range(height):
        t = min(1.0, y / (height * 0.78))
        color = tuple(int(top[i] * (1 - t) + bottom[i] * t) for i in range(3))
        draw.line((0, y, width, y), fill=color)
    nodes = [(70, 145), (220, 220), (380, 125), (545, 245), (730, 135), (165, 430), (415, 380), (675, 480)]
    for a, b in zip(nodes, nodes[1:]):
        draw.line((*a, *b), fill=(78, 110, 155), width=1)
    for x, y in nodes:
        draw.ellipse((x - 4, y - 4, x + 4, y + 4), fill=(170, 194, 225))
    logo_font = _font(105, bold=True)
    sub_font = _font(20, bold=True)
    logo = "2COM"
    box = draw.textbbox((0, 0), logo, font=logo_font)
    draw.text(((width - (box[2] - box[0])) / 2, 400), logo, fill="white", font=logo_font)
    sub = "CONSULTING"
    sbox = draw.textbbox((0, 0), sub, font=sub_font)
    draw.text(((width - (sbox[2] - sbox[0])) / 2, 515), sub, fill=(215, 228, 245), font=sub_font)
    panel_y = 795
    draw.rounded_rectangle((80, panel_y, 747, 1110), radius=12, fill=(250, 251, 253))
    title_font = _font(23, bold=True)
    label_font = _font(13, bold=True)
    value_font = _font(13)
    red = (166, 36, 46)
    black = (20, 24, 30)
    draw.text((115, panel_y + 38), "DOCUMENTAÇÃO", fill=red, font=title_font)
    client = _text(review.get("client")) or "CLIENTE"
    resp = dict(review.get("responsibles") or {})
    rows = [
        ("CLIENTE:", client),
        ("RESPONSÁVEL INFRA:", f"{_text(resp.get('infra'))} - 2COM CONSULTING" if _text(resp.get('infra')) else ""),
        ("RESPONSÁVEL DBA:", f"{_text(resp.get('dba'))} - 2COM CONSULTING" if _text(resp.get('dba')) else ""),
        ("RESPONSÁVEL NOC:", f"{_text(resp.get('noc'))} - 2COM CONSULTING" if _text(resp.get('noc')) else ""),
        ("REVISÃO:", f"{_text(resp.get('review'))} - 2COM CONSULTING" if _text(resp.get('review')) else ""),
        ("REVISÃO NOC:", f"{_text(resp.get('review_noc'))} - 2COM CONSULTING" if _text(resp.get('review_noc')) else ""),
        ("DATA:", _text(review.get("date")) or datetime.now().strftime("%d/%m/%Y")),
    ]
    y = panel_y + 86
    for label, value in rows:
        draw.text((115, y), label, fill=black, font=label_font)
        draw.text((320, y), _fit(value, 385, value_font), fill=black, font=value_font)
        y += 30
    out = io.BytesIO()
    image.save(out, format="PNG", optimize=True)
    return out.getvalue()


def _set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = BLUE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def _style_table(table, *, alternating: bool = True, header: bool = False) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell)
            if header and r_index == 0:
                _set_cell_fill(cell, "FFFFFF")
            elif alternating and r_index % 2 == 1:
                _set_cell_fill(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
                    if header and r_index == 0:
                        run.font.bold = True


def _set_page_layout(section, *, cover: bool = False) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    if cover:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0)
    else:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.82)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def _apply_doc_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor.from_string(DARK)
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(11.5)
    doc.styles["Heading 3"].font.bold = True


def _doc_header_footer(section) -> None:
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.18)
    hp = section.header.paragraphs[0]
    hp.clear()
    left = hp.add_run("2COM")
    left.font.name = "Arial"; left.font.size = Pt(17); left.font.bold = True; left.font.color.rgb = RGBColor.from_string(NAVY)
    right = hp.add_run("                                      INFRA  •  CLOUD  •  DATA")
    right.font.name = "Arial"; right.font.size = Pt(8); right.font.bold = True; right.font.color.rgb = RGBColor.from_string(NAVY)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("2COM CONSULTING                                      Matriz Goiânia\nRua 5, Qd.C4, Lt.16-19, Sala Comercial nº 115 · Setor Oeste, Goiânia - GO · +55 (62) 3095-8100")
    run.font.size = Pt(7.2); run.font.name = "Arial"; run.font.color.rgb = RGBColor.from_string(NAVY)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr(); p_bdr = OxmlElement("w:pBdr"); border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single"); border.set(qn("w:sz"), "18" if level == 1 else "12"); border.set(qn("w:space"), "1"); border.set(qn("w:color"), NAVY if level == 1 else BLUE)
    p_bdr.append(border); p_pr.append(p_bdr)


def _label_value_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2); table.cell(0, 0).text = title; table.cell(0, 1).text = ""
    for i, (label, value) in enumerate(rows, start=1): table.cell(i, 0).text = label; table.cell(i, 1).text = _text(value)
    _style_table(table, alternating=True); doc.add_paragraph()


def _grid_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for col, header in enumerate(headers): table.cell(0, col).text = header
    for r, values in enumerate(rows, start=1):
        for c, value in enumerate(values): table.cell(r, c).text = _text(value)
    _style_table(table, alternating=True, header=True); doc.add_paragraph()


def _field(values: dict[str, str], key: str, default: str = "") -> str:
    return _text(values.get(key) or default)


def _host_title(host: dict[str, Any]) -> str:
    role = _text(host.get("role") or host.get("kind") or "Servidor")
    mapping = {"monitoring_local": "Servidor de Monitoramento", "monitoring": "Servidor de Monitoramento", "application": "Servidor de Aplicação", "database": "Servidor de Banco de Dados", "standby": "Servidor de Banco de Dados Redundância", "training": "Servidor de Banco de Dados Teste"}
    return mapping.get(role.casefold(), f"Servidor - {role}")


def _selected_hosts(review: dict[str, Any]) -> list[dict[str, Any]]:
    return [dict(item) for item in review.get("selected_hosts") or [] if isinstance(item, dict)]


def export_docx(review: dict[str, Any]) -> bytes:
    review = sanitize_n2_review(deepcopy(review)); values = _field_map(review); doc = Document(); _apply_doc_defaults(doc)
    cover_section = doc.sections[0]; _set_page_layout(cover_section, cover=True)
    p = doc.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(build_cover_png(review)), width=Inches(8.27), height=Inches(11.69))
    section = doc.add_section(WD_SECTION.NEW_PAGE); _set_page_layout(section); _doc_header_footer(section)
    _heading(doc, "Sumário", 1)
    top_sections = {"Introdução", "Infraestrutura", "Banco de Dados", "Política de Backup", "Monitoramento", "Considerações Finais"}
    for item in SUMMARY_SECTIONS:
        p = doc.add_paragraph(item); p.paragraph_format.left_indent = Inches(0 if item in top_sections else 0.18); p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()
    _heading(doc, "Introdução", 1)
    doc.add_paragraph("A 2COM CONSULTING é uma empresa de outsourcing em TI com foco em infraestrutura e banco de dados. Esta documentação registra as evidências coletadas e revisadas para o ambiente do cliente, mantendo credenciais e segredos fora do documento.")
    _heading(doc, "Infraestrutura", 1); _heading(doc, "Inventário", 2)
    doc.add_paragraph("Esta etapa representa os hosts selecionados pelo analista N2 para validação documental.")
    hosts = _selected_hosts(review); inventory = doc.add_table(rows=max(1, len(hosts)), cols=3)
    if not hosts:
        inventory.cell(0, 0).text = "☐"; inventory.cell(0, 1).text = "Nenhum host selecionado"; inventory.cell(0, 2).text = ""
    else:
        for i, host in enumerate(hosts): inventory.cell(i, 0).text = "☒"; inventory.cell(i, 1).text = _host_title(host); inventory.cell(i, 2).text = _text(host.get("host"))
    _style_table(inventory, alternating=False); doc.add_paragraph(); _heading(doc, "Informações de Infraestrutura do Inventário", 2)
    for host in hosts:
        fields = dict(host.get("fields") or {})
        _label_value_table(doc, f"{_host_title(host)} - Hardware", [("Servidor", _text(fields.get("server") or host.get("host"))), ("Address Ipv4", _text(fields.get("address_ipv4") or host.get("ip"))), ("Address VPN", _text(fields.get("address_vpn"))), ("Nome Do Host", _text(fields.get("hostname") or host.get("host"))), ("Processador", _text(fields.get("processor"))), ("Memória", _text(fields.get("memory"))), ("Armazenamento", _text(fields.get("storage"))), ("Sistema Operacional", _text(fields.get("os")))])
        if _text(host.get("collection_notes")): doc.add_paragraph("Evidências/observações: " + _text(host.get("collection_notes")))
    _heading(doc, "Banco de Dados", 1)
    doc.add_paragraph("Etapa responsável por informações relacionadas ao software de banco de dados. A exportação nunca inclui senhas.")
    _label_value_table(doc, "Servidor de Banco de Dados - Software Oracle", [("Versão do Software Oracle - RDBMS", _field(values, "rdbms_version")), ("Instâncias (Separada por Vírgula)", _field(values, "instances"))])
    _label_value_table(doc, "Servidor de Banco de Dados - Password", [("WINT - SYS", ""), ("WINT - SYSTEM", ""), ("WINT - SCHEMA", "")])
    _heading(doc, "Informações para Ativação na TOTVS", 2)
    _label_value_table(doc, "Servidor de Banco de Dados - WINT", [("SID", _field(values, "sid")), ("USERNAME", _field(values, "totvs_username")), ("NUMERO DE SÉRIE", _field(values, "serial_number"))])
    _heading(doc, "Estrutura SGDB", 2); doc.add_paragraph(_field(values, "sgdb_structure") or "Não confirmado durante a coleta.")
    _heading(doc, "TNSNAMES", 2); doc.add_paragraph(_field(values, "tnsnames") or "Não confirmado durante a coleta.")
    if _field(values, "database_notes"): doc.add_paragraph("Evidências de banco: " + _field(values, "database_notes"))
    _heading(doc, "Mapeamento do Winthor", 1)
    doc.add_paragraph("A 2Com Consulting utiliza compartilhamentos com permissões restritivas para proteger as rotinas do WinThor. Este documento registra apenas usuários técnicos, caminhos e origens autorizadas; senhas permanecem em branco.")
    _grid_table(doc, ["Usuário", "Senha", "Path"], [[_field(values, "winthor_user_admin"), "", _field(values, "winthor_admin_path")], [_field(values, "winthor_user_read"), "", _field(values, "winthor_read_path")]])
    if _field(values, "winthor_notes"): doc.add_paragraph(_field(values, "winthor_notes"))
    _heading(doc, "Política de Backup", 1); doc.add_paragraph(_field(values, "backup_strategy") or "Estratégia de backup ainda não confirmada na revisão N2.")
    _heading(doc, "Database Oracle", 2)
    if _field(values, "instances"): doc.add_paragraph("Instância(s): " + _field(values, "instances"))
    doc.add_paragraph("Backups Lógicos do Banco de Dados")
    _grid_table(doc, ["Frequência", "Horário Início", "Horário Conclusão", "Duração", "Tam. Backup", "Cópia para Redundância"], [[_field(values, "datapump_frequency"), _field(values, "datapump_start"), _field(values, "datapump_end"), _field(values, "datapump_duration"), _field(values, "datapump_size"), _field(values, "datapump_redundancy")]])
    doc.add_paragraph("Backup Físico do Banco de Dados Utilizando a Ferramenta RMAN")
    _grid_table(doc, ["Frequência", "Tipo", "Horário Início", "Horário Conclusão", "Duração", "Tam. Backup", "Cópia para Redundância"], [[_field(values, "rman_frequency"), _field(values, "rman_type"), _field(values, "rman_start"), _field(values, "rman_end"), _field(values, "rman_duration"), _field(values, "rman_size"), _field(values, "rman_redundancy")], [_field(values, "archives_frequency"), "Archive (Incr.)", "-", "-", _field(values, "archives_duration"), " - ".join(item for item in (_field(values, "archives_size_min"), _field(values, "archives_size_max")) if item), _field(values, "archives_redundancy")]])
    _heading(doc, "Backup do Sistema ERP", 2); doc.add_paragraph("Winthor")
    _grid_table(doc, ["Frequência", "Horário Início", "Horário Conclusão", "Duração", "Tam. Backup", "Cópia para Redundância"], [[_field(values, "winthor_backup_frequency"), _field(values, "winthor_backup_start"), _field(values, "winthor_backup_end"), _field(values, "winthor_backup_duration"), _field(values, "winthor_backup_size"), _field(values, "winthor_backup_redundancy")]])
    _heading(doc, "Métodos de Execução e Validações", 1); _heading(doc, "Backup Lógico", 2)
    doc.add_paragraph(_field(values, "logical_backup_method") or _field(values, "backup_execution_notes") or "Não confirmado durante a coleta.")
    _heading(doc, "Backup Físico", 2); doc.add_paragraph(_field(values, "physical_backup_method") or _field(values, "backup_execution_notes") or "Não confirmado durante a coleta.")
    _heading(doc, "Retenção de Backups Obsoletos", 1); doc.add_paragraph(_field(values, "retention_notes") or "A política de retenção precisa ser confirmada/revisada pelo analista N2.")
    _label_value_table(doc, "Disco Local", [("Caminho", _field(values, "local_backup_path")), ("RMAN", _field(values, "rman_local_dir")), ("Datapump", _field(values, "datapump_local_dir")), ("Limite disco Datapump", _field(values, "datapump_local_threshold")), ("Mínimo Datapump", _field(values, "datapump_local_min")), ("Máximo Datapump", _field(values, "datapump_local_max")), ("Winthor", _field(values, "winthor_local_dir"))])
    _label_value_table(doc, "Unidade de Redundância", [("Caminho", _field(values, "redundancy_backup_path")), ("RMAN", _field(values, "rman_redundancy_dir")), ("Datapump", _field(values, "datapump_redundancy_dir")), ("Winthor", _field(values, "winthor_redundancy_dir"))])
    _heading(doc, "Dados Redundância", 1)
    _grid_table(doc, ["Tipo", "Modelo/Protocolo", "Capacidade", "Compartilhamento", "Usuário", "Senha"], [[_field(values, "redundancy_type"), _field(values, "redundancy_model"), _field(values, "redundancy_capacity"), _field(values, "redundancy_share"), _field(values, "redundancy_user"), ""]])
    _heading(doc, "Unidade de Redundância em Nuvem", 2); doc.add_paragraph(_field(values, "cloud_redundancy") or "Não confirmado durante a coleta.")
    _heading(doc, "Replicação standby", 1); _heading(doc, "Database Oracle", 2); doc.add_paragraph(_field(values, "standby_db_sync") or "Não confirmado durante a coleta.")
    _heading(doc, "Sistema Winthor", 2); doc.add_paragraph(_field(values, "standby_winthor_sync") or "Não confirmado durante a coleta.")
    if _field(values, "redundancy_notes"): doc.add_paragraph("Observações finais: " + _field(values, "redundancy_notes"))
    _heading(doc, "Monitoramento", 1); doc.add_paragraph(_field(values, "monitoring_notes") or "O monitoramento implementado nos servidores e serviços do escopo permite identificar problemas de banco de dados, backups e sistema operacional.")
    _grid_table(doc, ["URL", "Usuário", "Senha"], [[_field(values, "monitoring_url"), _field(values, "monitoring_user"), ""]])
    doc.add_paragraph(f"Site Checkmk: {_field(values, 'monitoring_site', _text(review.get('site_id')))}")
    doc.add_paragraph(f"Endpoint: {_field(values, 'monitoring_endpoint')}")
    doc.add_paragraph(f"Hosts monitorados: {_field(values, 'monitoring_host_count')}")
    doc.add_paragraph(f"Problemas ativos no momento da coleta: {_field(values, 'monitoring_problem_count')}")
    _heading(doc, "Considerações Finais", 1); doc.add_paragraph(_field(values, "closing_notes") or "Documentação gerada a partir das evidências coletadas e revisadas pelo analista N2.")
    doc.add_paragraph("Campos sem evidência permanecem em branco ou explicitamente marcados como não confirmados. Credenciais não são exportadas.")
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle("N2H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=19, textColor=colors.HexColor("#10151D"), spaceBefore=8, spaceAfter=7),
        "h2": ParagraphStyle("N2H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#10151D"), spaceBefore=7, spaceAfter=5),
        "body": ParagraphStyle("N2Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12, textColor=colors.HexColor("#10151D"), spaceAfter=6, alignment=TA_LEFT),
        "small": ParagraphStyle("N2Small", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.5, leading=10, textColor=colors.HexColor("#5A6472"), spaceAfter=4),
    }


def _pdf_header_footer(c: canvas.Canvas, doc) -> None:
    width, height = A4; c.saveState(); c.setFillColor(colors.HexColor("#0B1F44")); c.setFont("Helvetica-Bold", 16); c.drawString(16 * mm, height - 16 * mm, "2COM")
    c.setFont("Helvetica-Bold", 7); c.drawRightString(width - 16 * mm, height - 15 * mm, "INFRA  •  CLOUD  •  DATA")
    c.setStrokeColor(colors.HexColor("#2672D5")); c.setLineWidth(0.7); c.line(16 * mm, height - 19 * mm, width - 16 * mm, height - 19 * mm)
    c.setFont("Helvetica-Bold", 7.2); c.drawString(16 * mm, 15 * mm, "2COM CONSULTING"); c.drawRightString(width - 16 * mm, 19 * mm, "Matriz Goiânia")
    c.setFont("Helvetica", 6.4); c.drawRightString(width - 16 * mm, 15 * mm, "Rua 5, Qd.C4, Lt.16-19 · Setor Oeste, Goiânia - GO"); c.drawRightString(width - 16 * mm, 11.5 * mm, "+55 (62) 3095-8100"); c.restoreState()


def _pdf_table(headers: list[str], rows: list[list[str]], widths=None):
    table = Table([headers] + rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTNAME", (0, 1), (-1, -1), "Helvetica"), ("FONTSIZE", (0, 0), (-1, -1), 7.6), ("LEADING", (0, 0), (-1, -1), 9.5), ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#10151D")), ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#2672D5")), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#DCE9FC")]), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)])); return table


def _pdf_label_table(title: str, rows: list[tuple[str, str]], styles):
    body = [[Paragraph(f"<b>{title}</b>", styles["small"]), ""]]; body.extend([[Paragraph(_text(label), styles["small"]), Paragraph(_text(value), styles["small"])] for label, value in rows])
    table = Table(body, colWidths=[85 * mm, 85 * mm], hAlign="LEFT")
    table.setStyle(TableStyle([("SPAN", (0, 0), (1, 0)), ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#2672D5")), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#DCE9FC"), colors.white]), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)])); return table


def _content_pdf(review: dict[str, Any]) -> bytes:
    values = _field_map(review); styles = _pdf_styles(); out = io.BytesIO(); doc = SimpleDocTemplate(out, pagesize=A4, rightMargin=16 * mm, leftMargin=16 * mm, topMargin=28 * mm, bottomMargin=28 * mm)
    story: list[Any] = []; h1, h2, body, small = styles["h1"], styles["h2"], styles["body"], styles["small"]
    story.append(Paragraph("Sumário", h1)); [story.append(Paragraph(item, body)) for item in SUMMARY_SECTIONS]; story.append(PageBreak())
    story += [Paragraph("Introdução", h1), Paragraph("A 2COM CONSULTING é uma empresa de outsourcing em TI com foco em infraestrutura e banco de dados. Esta documentação registra as evidências coletadas e revisadas para o ambiente do cliente, mantendo credenciais e segredos fora do documento.", body), Paragraph("Infraestrutura", h1), Paragraph("Inventário", h2), Paragraph("Esta etapa representa os hosts selecionados pelo analista N2 para validação documental.", body)]
    hosts = _selected_hosts(review); story.append(_pdf_table(["Tipo", "Host", "IP"], [[_host_title(h), _text(h.get("host")), _text((h.get("fields") or {}).get("address_ipv4") or h.get("ip"))] for h in hosts] or [["-", "Nenhum host selecionado", "-"]], [60*mm, 65*mm, 45*mm])); story += [Spacer(1, 5*mm), Paragraph("Informações de Infraestrutura do Inventário", h2)]
    for host in hosts:
        fields = dict(host.get("fields") or {}); story.append(_pdf_label_table(f"{_host_title(host)} - Hardware", [("Servidor", _text(fields.get("server") or host.get("host"))), ("Address Ipv4", _text(fields.get("address_ipv4") or host.get("ip"))), ("Address VPN", _text(fields.get("address_vpn"))), ("Nome Do Host", _text(fields.get("hostname") or host.get("host"))), ("Processador", _text(fields.get("processor"))), ("Memória", _text(fields.get("memory"))), ("Armazenamento", _text(fields.get("storage"))), ("Sistema Operacional", _text(fields.get("os")))], styles))
        if _text(host.get("collection_notes")): story.append(Paragraph("Evidências/observações: " + _text(host.get("collection_notes")), small))
        story.append(Spacer(1, 3*mm))
    story += [Paragraph("Banco de Dados", h1), Paragraph("Etapa responsável por informações relacionadas ao software de banco de dados. A exportação nunca inclui senhas.", body)]
    story.append(_pdf_label_table("Servidor de Banco de Dados - Software Oracle", [("Versão do Software Oracle - RDBMS", _field(values,"rdbms_version")), ("Instâncias", _field(values,"instances"))], styles)); story += [Spacer(1, 3*mm), Paragraph("Informações para Ativação na TOTVS", h2)]
    story.append(_pdf_label_table("Servidor de Banco de Dados - WINT", [("SID", _field(values,"sid")), ("USERNAME", _field(values,"totvs_username")), ("NUMERO DE SÉRIE", _field(values,"serial_number"))], styles)); story += [Spacer(1, 3*mm), Paragraph("Estrutura SGDB", h2), Paragraph(_field(values,"sgdb_structure") or "Não confirmado durante a coleta.", body), Paragraph("TNSNAMES", h2), Paragraph(_field(values,"tnsnames") or "Não confirmado durante a coleta.", body)]
    if _field(values,"database_notes"): story.append(Paragraph("Evidências de banco: "+_field(values,"database_notes"), small))
    story += [Paragraph("Mapeamento do Winthor", h1), Paragraph("São registrados somente usuários técnicos, caminhos e origens autorizadas; senhas permanecem em branco.", body)]; story.append(_pdf_table(["Usuário","Senha","Path"], [[_field(values,"winthor_user_admin"),"",_field(values,"winthor_admin_path")],[_field(values,"winthor_user_read"),"",_field(values,"winthor_read_path")]], [45*mm,35*mm,90*mm]))
    if _field(values,"winthor_notes"): story.append(Paragraph(_field(values,"winthor_notes"), small))
    story += [Paragraph("Política de Backup", h1), Paragraph(_field(values,"backup_strategy") or "Estratégia de backup ainda não confirmada na revisão N2.", body), Paragraph("Database Oracle", h2), Paragraph("Backups Lógicos do Banco de Dados", body)]
    story.append(_pdf_table(["Frequência","Início","Conclusão","Duração","Tam.","Redundância"], [[_field(values,"datapump_frequency"),_field(values,"datapump_start"),_field(values,"datapump_end"),_field(values,"datapump_duration"),_field(values,"datapump_size"),_field(values,"datapump_redundancy")]], [28*mm,25*mm,28*mm,25*mm,25*mm,39*mm])); story += [Spacer(1,3*mm), Paragraph("Backup Físico do Banco de Dados Utilizando a Ferramenta RMAN", body)]
    story.append(_pdf_table(["Frequência","Tipo","Início","Conclusão","Duração","Tam.","Redundância"], [[_field(values,"rman_frequency"),_field(values,"rman_type"),_field(values,"rman_start"),_field(values,"rman_end"),_field(values,"rman_duration"),_field(values,"rman_size"),_field(values,"rman_redundancy")],[_field(values,"archives_frequency"),"Archive (Incr.)","-","-",_field(values,"archives_duration")," - ".join(v for v in (_field(values,"archives_size_min"),_field(values,"archives_size_max")) if v),_field(values,"archives_redundancy")]], [24*mm,28*mm,22*mm,25*mm,22*mm,22*mm,27*mm])); story += [Paragraph("Backup do Sistema ERP", h2), Paragraph("Winthor", body)]
    story.append(_pdf_table(["Frequência","Início","Conclusão","Duração","Tam.","Redundância"], [[_field(values,"winthor_backup_frequency"),_field(values,"winthor_backup_start"),_field(values,"winthor_backup_end"),_field(values,"winthor_backup_duration"),_field(values,"winthor_backup_size"),_field(values,"winthor_backup_redundancy")]], [28*mm,25*mm,28*mm,25*mm,25*mm,39*mm]))
    story += [Paragraph("Métodos de Execução e Validações", h1), Paragraph("Backup Lógico", h2), Paragraph(_field(values,"logical_backup_method") or _field(values,"backup_execution_notes") or "Não confirmado durante a coleta.", body), Paragraph("Backup Físico", h2), Paragraph(_field(values,"physical_backup_method") or _field(values,"backup_execution_notes") or "Não confirmado durante a coleta.", body), Paragraph("Retenção de Backups Obsoletos", h1), Paragraph(_field(values,"retention_notes") or "A política de retenção precisa ser confirmada/revisada pelo analista N2.", body)]
    story.append(_pdf_label_table("Disco Local", [("Caminho",_field(values,"local_backup_path")),("RMAN",_field(values,"rman_local_dir")),("Datapump",_field(values,"datapump_local_dir")),("Limite Datapump",_field(values,"datapump_local_threshold")),("Mínimo Datapump",_field(values,"datapump_local_min")),("Máximo Datapump",_field(values,"datapump_local_max")),("Winthor",_field(values,"winthor_local_dir"))], styles)); story.append(Spacer(1,3*mm)); story.append(_pdf_label_table("Unidade de Redundância", [("Caminho",_field(values,"redundancy_backup_path")),("RMAN",_field(values,"rman_redundancy_dir")),("Datapump",_field(values,"datapump_redundancy_dir")),("Winthor",_field(values,"winthor_redundancy_dir"))], styles))
    story += [Paragraph("Dados Redundância", h1)]; story.append(_pdf_table(["Tipo","Modelo/Protocolo","Capacidade","Compartilhamento","Usuário","Senha"], [[_field(values,"redundancy_type"),_field(values,"redundancy_model"),_field(values,"redundancy_capacity"),_field(values,"redundancy_share"),_field(values,"redundancy_user"),""]], [22*mm,35*mm,27*mm,40*mm,28*mm,18*mm]))
    story += [Paragraph("Unidade de Redundância em Nuvem", h2), Paragraph(_field(values,"cloud_redundancy") or "Não confirmado durante a coleta.", body), Paragraph("Replicação standby", h1), Paragraph("Database Oracle", h2), Paragraph(_field(values,"standby_db_sync") or "Não confirmado durante a coleta.", body), Paragraph("Sistema Winthor", h2), Paragraph(_field(values,"standby_winthor_sync") or "Não confirmado durante a coleta.", body)]
    if _field(values,"redundancy_notes"): story.append(Paragraph("Observações finais: "+_field(values,"redundancy_notes"), body))
    story += [Paragraph("Monitoramento", h1), Paragraph(_field(values,"monitoring_notes") or "O monitoramento implementado nos servidores e serviços do escopo permite identificar problemas de banco de dados, backups e sistema operacional.", body)]; story.append(_pdf_table(["URL","Usuário","Senha"], [[_field(values,"monitoring_url"),_field(values,"monitoring_user"),""]], [80*mm,60*mm,30*mm])); story.append(Paragraph(f"Site Checkmk: {_field(values,'monitoring_site',_text(review.get('site_id')))} · Endpoint: {_field(values,'monitoring_endpoint')} · Hosts: {_field(values,'monitoring_host_count')} · Problemas ativos: {_field(values,'monitoring_problem_count')}", small))
    story += [Paragraph("Considerações Finais", h1), Paragraph(_field(values,"closing_notes") or "Documentação gerada a partir das evidências coletadas e revisadas pelo analista N2.", body), Paragraph("Campos sem evidência permanecem em branco ou explicitamente marcados como não confirmados. Credenciais não são exportadas.", small)]
    doc.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer); return out.getvalue()


def export_pdf(review: dict[str, Any]) -> bytes:
    review = sanitize_n2_review(deepcopy(review)); cover = io.BytesIO(); c = canvas.Canvas(cover, pagesize=A4)
    c.drawImage(ImageReader(io.BytesIO(build_cover_png(review))), 0, 0, width=A4[0], height=A4[1], preserveAspectRatio=False, mask="auto"); c.showPage(); c.save(); content = _content_pdf(review); writer = PdfWriter()
    for source in (cover.getvalue(), content):
        reader = PdfReader(io.BytesIO(source))
        for page in reader.pages: writer.add_page(page)
    out = io.BytesIO(); writer.write(out); return out.getvalue()


def export_n2_document(review: dict[str, Any], document_format: str) -> tuple[bytes, str, str]:
    fmt = _text(document_format).lower(); client = _clean_filename(_text(review.get("client")) or "cliente"); stamp = datetime.now().strftime("%Y%m%d")
    if fmt == "docx": return export_docx(review), f"Documentacao-N2-{client}-{stamp}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf": return export_pdf(review), f"Documentacao-N2-{client}-{stamp}.pdf", "application/pdf"
    raise ValueError("formato de exportação inválido; use docx ou pdf")
