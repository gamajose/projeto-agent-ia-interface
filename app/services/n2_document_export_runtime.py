from __future__ import annotations

import io
from typing import Any

from docx import Document as OpenDocument
from docx.shared import Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

from app.services import n2_document_export as _export


# python-docx cria um documento novo sem parágrafos. O gerador do N2 usa a
# primeira página como uma imagem de capa em um parágrafo de largura total; este
# factory garante que esse parágrafo exista sem alterar o restante do layout.
_OriginalDocument = _export.Document


def _document_with_cover_paragraph(*args: Any, **kwargs: Any):
    document = _OriginalDocument(*args, **kwargs)
    if not document.paragraphs:
        document.add_paragraph()
    return document


_export.Document = _document_with_cover_paragraph


def _client_name(review: dict[str, Any]) -> str:
    return str(review.get("client") or "").strip()


def _docx_with_searchable_client(content: bytes, review: dict[str, Any]) -> bytes:
    client = _client_name(review)
    if not client:
        return content
    document = OpenDocument(io.BytesIO(content))
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Cliente: {client}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(11, 31, 68)

    # Move o identificador logo após o título Introdução. A capa continua sendo
    # a referência visual principal, mas o cliente também passa a existir como
    # texto pesquisável no XML do Word.
    for item in document.paragraphs:
        if item is paragraph:
            continue
        if item.text.strip().casefold() == "introdução":
            paragraph._p.getparent().remove(paragraph._p)
            item._p.addnext(paragraph._p)
            break

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_with_searchable_client(content: bytes, review: dict[str, Any]) -> bytes:
    client = _client_name(review)
    if not client:
        return content
    reader = PdfReader(io.BytesIO(content))
    writer = PdfWriter()
    for index, page in enumerate(reader.pages):
        # A página 0 é a capa gráfica. Na primeira página de conteúdo, registra
        # o cliente na área entre o cabeçalho e o conteúdo para manter pesquisa
        # e acessibilidade sem interferir no layout do template.
        if index == 1:
            packet = io.BytesIO()
            overlay = canvas.Canvas(packet, pagesize=A4)
            overlay.setFillColor(colors.HexColor("#0B1F44"))
            overlay.setFont("Helvetica-Bold", 7.2)
            overlay.drawString(16 * mm, A4[1] - 23 * mm, f"Cliente: {client}")
            overlay.save()
            packet.seek(0)
            overlay_page = PdfReader(packet).pages[0]
            page.merge_page(overlay_page)
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def export_docx(review: dict[str, Any]) -> bytes:
    return _docx_with_searchable_client(_export.export_docx(review), review)


def export_pdf(review: dict[str, Any]) -> bytes:
    return _pdf_with_searchable_client(_export.export_pdf(review), review)


def export_n2_document(review: dict[str, Any], document_format: str) -> tuple[bytes, str, str]:
    content, filename, media_type = _export.export_n2_document(review, document_format)
    fmt = str(document_format or "").strip().lower()
    if fmt == "docx":
        content = _docx_with_searchable_client(content, review)
    elif fmt == "pdf":
        content = _pdf_with_searchable_client(content, review)
    return content, filename, media_type


build_cover_png = _export.build_cover_png
